"""
Comprehensive DOCX Report Generator for ANOVA Analysis
======================================================
Generates fully editable Word (.docx) reports with all tables and graphics.
Mirrors the PDF report content so users can make corrections after export.
"""

import io
import numpy as np
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── colour helpers ────────────────────────────────────────────────────────────

def _rgb(hex_str: str) -> RGBColor:
    """Convert '#RRGGBB' string to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade_cell(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    """Set individual column widths in a Word table."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ── paragraph / heading helpers ───────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int, color_hex: str = "#1f77b4"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = _rgb(color_hex)
    return p


def _normal(doc: Document, text: str, italic: bool = False, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


# ── table builder ─────────────────────────────────────────────────────────────

def _build_table(doc: Document, data: list, col_widths_cm: list,
                 header_hex: str = "#4472C4", font_size: int = 9,
                 alternate: bool = True):
    """
    Build a Word table from a list-of-lists.
    data[0] = header row.
    """
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0] if para.runs else para.add_run(str(cell_text))
            run.font.size = Pt(font_size)

            if r_idx == 0:                          # header
                _shade_cell(cell, header_hex)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif alternate and r_idx % 2 == 0:     # even body rows
                _shade_cell(cell, "E9E9E9")

    _set_col_widths(table, col_widths_cm)
    return table


# ── main public function ──────────────────────────────────────────────────────

def generate_comprehensive_docx_report(
    config,
    detailed_anova_df,
    model_summary,
    model_results,
    significant_terms,
    sig_terms_df,
    variance_display_df,
    half_normal_fig=None,
    residuals_fig=None,
    histogram_fig=None,
    predicted_fig=None,
    initial_data_df=None,
) -> bytes:
    """
    Generate a comprehensive, editable DOCX report.

    Parameters mirror :func:`generate_comprehensive_pdf_report` exactly so the
    Streamlit app can call both with the same arguments.

    Returns
    -------
    bytes
        Raw .docx file content ready for st.download_button.
    """

    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading("COMPREHENSIVE ANOVA ANALYSIS REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _rgb("#1f77b4")
        run.font.size = Pt(20)

    ts = _normal(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", italic=True)
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # =========================================================================
    # 1. ANALYSIS CONFIGURATION
    # =========================================================================
    _heading(doc, "1. Analysis Configuration", 1, "#ff7f0e")

    info_data = [
        ["Parameter", "Value"],
        ["Analysis Type",       config.get("analysis_type", "N/A")],
        ["Model Type",          config.get("model_type", "N/A").title()],
        ["Response Variable",   config.get("response_col", "N/A")],
        ["Number of Observations", str(len(model_results["residuals"]))],
        ["Number of Factors",   str(len(config.get("factor_cols", [])))],
        ["Factor Names",        ", ".join(config.get("factor_cols", [])[:6])
                                + ("…" if len(config.get("factor_cols", [])) > 6 else "")],
    ]
    _build_table(doc, info_data, [7, 9], header_hex="808080", font_size=10)
    doc.add_paragraph()

    # =========================================================================
    # 2. MODEL PERFORMANCE SUMMARY
    # =========================================================================
    _heading(doc, "2. Model Performance Summary", 1, "#ff7f0e")

    r2 = model_summary.get("R_squared", float("nan"))
    perf_data = [
        ["Metric", "Value", "Interpretation"],
        ["R² Score",        f"{r2:.4f}",
         "Excellent" if r2 > 0.9 else ("Good" if r2 > 0.7 else "Fair")],
        ["Adjusted R²",     f"{model_summary.get('R_squared_adj', float('nan')):.4f}",
         "Accounts for model complexity"],
        ["RMSE",            f"{model_results.get('rmse', float('nan')):.4f}",
         "Root Mean Square Error"],
        ["Residual MS",     f"{model_summary.get('MS_error', float('nan')):.4f}",
         "Mean square error"],
        ["Error DF",        str(int(model_summary.get("df_error", 0))),
         "Degrees of freedom"],
        ["Total Parameters", str(model_summary.get("n_parameters", "N/A")),
         "Model complexity"],
        ["Sample Size",     str(model_summary.get("n_samples", "N/A")),
         "Number of observations"],
    ]
    _build_table(doc, perf_data, [5, 4, 7], header_hex="4472C4", font_size=10)
    doc.add_paragraph()

    # =========================================================================
    # 3. SIGNIFICANT TERMS
    # =========================================================================
    if significant_terms:
        _heading(doc, f"3. Significant Terms  ({len(significant_terms)} found at α = 0.05)", 1, "#ff7f0e")
        shown = significant_terms[:30]
        for term in shown:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(term)).bold = True
        if len(significant_terms) > 30:
            _normal(doc, f"… and {len(significant_terms) - 30} more terms.", italic=True)
        doc.add_paragraph()

    # =========================================================================
    # 4. INITIAL EXPERIMENTAL DATA
    # =========================================================================
    if initial_data_df is not None and len(initial_data_df) > 0:
        doc.add_page_break()
        _heading(doc, "4. Initial Experimental Data", 1, "#ff7f0e")
        _normal(doc, "Design matrix and response values used for analysis.", italic=True)
        doc.add_paragraph()

        cols = list(initial_data_df.columns)
        max_per_chunk = 9          # columns per sub-table to fit page width

        for col_start in range(0, len(cols), max_per_chunk):
            chunk_cols = cols[col_start: col_start + max_per_chunk]
            header = chunk_cols
            rows = [header]
            for idx in range(len(initial_data_df)):
                row = []
                for c in chunk_cols:
                    val = initial_data_df.iloc[idx][c]
                    if isinstance(val, (int, np.integer)):
                        row.append(str(int(val)))
                    elif isinstance(val, (float, np.floating)):
                        row.append(f"{val:.3f}")
                    else:
                        row.append(str(val)[:20])
                rows.append(row)

            avail = 16.0            # usable width in cm
            cw = avail / len(chunk_cols)
            _build_table(doc, rows, [cw] * len(chunk_cols),
                         header_hex="2E75B6", font_size=8)
            doc.add_paragraph()

            if col_start + max_per_chunk < len(cols):
                _normal(doc, f"(Continued — columns {col_start + 1}–"
                             f"{min(col_start + max_per_chunk, len(cols))} of {len(cols)})",
                        italic=True)

        _normal(doc, f"Complete table: {len(initial_data_df)} observations.", italic=True)

    # =========================================================================
    # 5. DETAILED ANOVA TABLE
    # =========================================================================
    doc.add_page_break()
    _heading(doc, "5. Detailed ANOVA Table", 1, "#ff7f0e")
    _normal(doc, "Complete analysis of variance for all model terms.", italic=True)
    doc.add_paragraph()

    has_pseudo = "Pseudo_t_Ratio" in detailed_anova_df.columns

    if has_pseudo:
        anova_header = ["Source", "SS", "DF", "MS", "F", "P-Value",
                        "Pseudo-t", "Pseudo-P", "Coef"]
        col_w = [4.0, 1.8, 0.9, 1.8, 1.5, 1.8, 1.4, 1.8, 1.6]
    else:
        anova_header = ["Source", "SS", "DF", "MS", "F", "P-Value", "Coef"]
        col_w = [4.5, 2.2, 1.0, 2.2, 1.8, 2.2, 2.1]

    anova_rows = [anova_header]
    for _, row in detailed_anova_df.iterrows():
        f_str = ""
        if not pd.isna(row["F_statistic"]):
            f_str = (f"{row['F_statistic']:.2e}" if abs(row["F_statistic"]) > 1000
                     else f"{row['F_statistic']:.2f}")

        r = [
            str(row["Source"])[:35],
            f"{row['Sum_of_Squares']:.3f}" if not pd.isna(row["Sum_of_Squares"]) else "",
            f"{int(row['DF'])}"            if not pd.isna(row["DF"]) and row["DF"] >= 0 else "",
            f"{row['Mean_Square']:.3f}"    if not pd.isna(row["Mean_Square"]) else "",
            f_str,
            f"{row['P_Value']:.2e}"        if not pd.isna(row["P_Value"]) else "",
        ]
        if has_pseudo:
            r += [
                f"{row['Pseudo_t_Ratio']:.2f}"  if not pd.isna(row["Pseudo_t_Ratio"]) else "",
                f"{row['Pseudo_P_Value']:.2e}"  if not pd.isna(row["Pseudo_P_Value"]) else "",
            ]
        r.append(f"{row['Coefficient']:.3f}" if not pd.isna(row["Coefficient"]) else "")
        anova_rows.append(r)

    # split into 30-row chunks so table doesn't overflow
    chunk = 30
    for i in range(0, len(anova_rows) - 1, chunk):
        _build_table(doc, [anova_header] + anova_rows[1 + i: 1 + i + chunk],
                     col_w, header_hex="70AD47", font_size=8)
        if i + chunk < len(anova_rows) - 1:
            _normal(doc, f"(rows {i + 1}–{min(i + chunk, len(anova_rows) - 1)} "
                         f"of {len(anova_rows) - 1}, continued…)", italic=True)
            doc.add_paragraph()

    doc.add_paragraph()

    # =========================================================================
    # 6. VARIANCE ANALYSIS
    # =========================================================================
    if variance_display_df is not None and len(variance_display_df) > 0:
        doc.add_page_break()
        _heading(doc, "6. Variance Analysis for Effect Estimates", 1, "#ff7f0e")
        _normal(doc, "Standard errors and 95% confidence intervals for model coefficients.",
                italic=True)
        doc.add_paragraph()

        var_header = ["Term", "Coefficient", "Std Error",
                      "t-Stat", "95% CI Lower", "95% CI Upper", "Sig?"]
        var_rows = [var_header]
        for _, row in variance_display_df.iterrows():
            t_str = ""
            if not pd.isna(row["t-Stat"]):
                t_str = (f"{row['t-Stat']:.2e}" if abs(row["t-Stat"]) > 1000
                         else f"{row['t-Stat']:.2f}")
            var_rows.append([
                str(row["Term"])[:28],
                f"{row['Coefficient']:.4f}",
                f"{row['Std Error']:.4f}",
                t_str,
                f"{row['95% CI Lower']:.4f}",
                f"{row['95% CI Upper']:.4f}",
                "Yes" if row["Significant"] else "No",
            ])

        _build_table(doc, var_rows, [4.5, 2.2, 2.2, 1.8, 2.4, 2.4, 1.0],
                     header_hex="5B9BD5", font_size=8)
        doc.add_paragraph()

    # =========================================================================
    # 7. GRAPHICAL ANALYSIS  (embedded as PNG images)
    # =========================================================================
    figs = [
        (half_normal_fig,  "7.1 Half-Normal Probability Plot",
         "Identifies significant effects graphically."),
        (residuals_fig,    "7.2 Residuals vs Fitted Values",
         "Check for patterns in residuals."),
        (histogram_fig,    "7.3 Residuals Distribution",
         "Check for normality of residuals."),
        (predicted_fig,    "7.4 Actual vs Predicted Values",
         "Model prediction accuracy."),
    ]

    any_fig = any(f[0] is not None for f in figs)
    if any_fig:
        doc.add_page_break()
        _heading(doc, "7. Graphical Analysis", 1, "#ff7f0e")
        doc.add_paragraph()

    for fig_obj, fig_title, fig_desc in figs:
        if fig_obj is None:
            continue
        try:
            _heading(doc, fig_title, 2, "#2ca02c")
            _normal(doc, fig_desc, italic=True)
            img_bytes = fig_obj.to_image(format="png", width=800, height=500, scale=2)
            doc.add_picture(io.BytesIO(img_bytes), width=Cm(15))
            doc.add_paragraph()
        except Exception as e:
            _normal(doc, f"[Figure could not be embedded: {e}]", italic=True)

    # =========================================================================
    # 8. FINAL MODEL FORMULA
    # =========================================================================
    if sig_terms_df is not None and len(sig_terms_df) > 0:
        doc.add_page_break()
        _heading(doc, "8. Final Model Formula", 1, "#ff7f0e")
        _normal(doc, f"Fitted model with {len(sig_terms_df)} significant terms.")
        doc.add_paragraph()

        # Build formula string
        parts = []
        for _, row in sig_terms_df.iterrows():
            term = row["Term"]
            coef = row["Coefficient"]
            sign = "+ " if (coef >= 0 and parts) else ("" if not parts else "- ")
            parts.append(f"{sign}{abs(coef):.4f} × {term}")

        formula_str = "Response  =  " + "  ".join(parts)

        # Wrap long formula across lines (every ~90 chars)
        line_width = 90
        words = formula_str.split("  ")
        lines_ = []
        cur = ""
        for w in words:
            if len(cur) + len(w) + 2 > line_width and cur:
                lines_.append(cur)
                cur = "            " + w          # indent continuation
            else:
                cur = (cur + "  " + w).strip() if cur else w
        if cur:
            lines_.append(cur)

        p = doc.add_paragraph()
        p.style = "No Spacing"
        for ln in lines_:
            run = p.add_run(ln + "\n")
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        doc.add_paragraph()

        # Summary metrics table
        _heading(doc, "Summary Metrics", 2, "#2ca02c")
        summary_data = [
            ["Metric", "Value"],
            ["Terms in Final Model", str(len(sig_terms_df))],
            ["R² Score",            f"{model_summary.get('R_squared', float('nan')):.4f}"],
            ["Adjusted R²",         f"{model_summary.get('R_squared_adj', float('nan')):.4f}"],
            ["RMSE",                f"{model_results.get('rmse', float('nan')):.4f}"],
        ]
        _build_table(doc, summary_data, [8, 6], header_hex="ED7D31", font_size=11)

    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()
    _normal(doc, "— End of Report —", italic=True)
    _normal(doc,
            f"Generated by DOE Analysis System  —  "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            italic=True)

    # ── serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
