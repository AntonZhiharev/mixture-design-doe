# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
"""Iteration 59 / ASSISTANT_SPEC — ВЛОЖЕНИЯ сессии ассистента.

Реальный вход технолога — документ: паспорт сырья (PDF), выгрузка лаборатории
(xlsx/csv), протокол (docx), кусок спеки (json). Ассистенту ЗАПРЕЩЕНО выдумывать
данные паспортов (d50, topcut, плотность), поэтому слой вложений обязан либо
дать текст, либо ЯВНО сказать, почему текста нет: молча пустое вложение
читается как «в паспорте ничего нет» — это худший отказ из возможных.

Покрытие: sha256-дедуп (файл не занимает контекст дважды); извлечение текста
из txt/csv/json/xlsx/docx/pdf; PDF-скан без текстового слоя → примечание про
OCR, а не тишина; отказ по формату/размеру/пустому файлу; усечение длинного
текста с дочитыванием хвоста с диска; хранение файла В ПРОЕКТЕ и удаление;
дайджест для контекста.
"""
import io
import json

import pytest

from src.assistant import files as af
from src.assistant import store, views
from src.assistant.session import new_session

PROJECT = "pvc_edge_v1"


def _session():
    return new_session(PROJECT, model="anthropic/claude-sonnet-4.5")


# ---------------------------------------------------------------- фикстуры
def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Лаборатория"
    ws.append(["опыт", "Gloss", "Adhesion"])
    ws.append([1, 8.4, 3.2])
    ws.append([2, 11.7, 2.9])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes() -> bytes:
    import docx
    doc = docx.Document()
    doc.add_paragraph("Протокол испытаний партии 2026-07")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "показатель"
    t.cell(0, 1).text = "значение"
    t.cell(1, 0).text = "ΔE"
    t.cell(1, 1).text = "0.8"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_bytes(text: str = "") -> bytes:
    """PDF с текстовым слоем (через reportlab) или пустая страница (скан)."""
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    if text:
        c.drawString(72, 720, text)
    c.showPage()
    c.save()
    return buf.getvalue()


# ----------------------------------------------------------------------
# 1. Дедуп и хранение в проекте
# ----------------------------------------------------------------------
def test_attach_stores_file_in_project_dir(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "TDS_Chalk_1T.txt",
                         "Мел молотый. d50 не указан.".encode("utf-8"))
    path = af.attachment_path(tmp_path, PROJECT, att)
    assert path.exists() and path.parent.name == "files"
    assert path.parent.parent.name == "assistant"
    assert att.sha256[:16] in path.name and att.mime == "text/plain"
    assert "d50 не указан" in att.text


def test_same_content_deduped_even_under_other_name(tmp_path):
    """Тот же файл под другим именем не должен занимать контекст дважды."""
    s = _session()
    data = b"identical payload"
    a1 = af.attach_file(s, tmp_path, "tds.txt", data)
    a2 = af.attach_file(s, tmp_path, "tds_copy.txt", data)
    assert a2 is a1 and len(s.attachments) == 1
    assert len(list((tmp_path / PROJECT / "assistant" / "files").iterdir())) == 1


def test_two_different_files_kept_separately(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "a.txt", b"aaa")
    af.attach_file(s, tmp_path, "b.txt", b"bbb")
    assert len(s.attachments) == 2


def test_attach_requires_project(tmp_path):
    """Сессия без проекта — некуда класть файл: явный отказ (A0.6)."""
    s = new_session("")
    with pytest.raises(af.AttachmentError, match="не привязана к проекту"):
        af.attach_file(s, tmp_path, "a.txt", b"data")


# ----------------------------------------------------------------------
# 2. Извлечение текста по форматам
# ----------------------------------------------------------------------
def test_extract_csv_and_json(tmp_path):
    s = _session()
    csv = af.attach_file(s, tmp_path, "runs.csv",
                         b"opyt,gloss\n1,8.4\n2,11.7\n")
    assert "gloss" in csv.text and csv.mime == "text/csv"

    js = af.attach_file(s, tmp_path, "spec.json",
                        json.dumps({"node": "UV_CSFCP",
                                    "role": "ABSOLUTE_CAPPED"},
                                   ensure_ascii=False).encode("utf-8"))
    assert "ABSOLUTE_CAPPED" in js.text
    assert js.text.count("\n") >= 2, "JSON должен быть развёрнут для чтения"


def test_broken_json_attached_as_text_with_note(tmp_path):
    """Битый JSON — не отказ: текст всё равно полезен, но причина видна."""
    s = _session()
    att = af.attach_file(s, tmp_path, "spec.json", "{не json".encode("utf-8"))

    assert "не json" in att.text and "JSON не разобран" in att.note


def test_extract_xlsx_sheets_and_cells(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "lab.xlsx", _xlsx_bytes())
    assert "### Лист: Лаборатория" in att.text
    assert "Adhesion" in att.text and "11.7" in att.text


def test_extract_docx_paragraphs_and_tables(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "protocol.docx", _docx_bytes())
    assert "Протокол испытаний" in att.text
    assert "### Таблица 1" in att.text and "0.8" in att.text


def test_extract_pdf_with_text_layer(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "TDS.pdf",
                         _pdf_bytes("Chalk 1T: d50 = 2.5 um"))
    assert "d50" in att.text and "— страница 1 —" in att.text
    assert att.mime == "application/pdf"


def test_pdf_scan_without_text_layer_explains_itself(tmp_path):
    """Скан без текста → примечание про OCR, а НЕ молча пустое вложение."""
    s = _session()
    att = af.attach_file(s, tmp_path, "scan.pdf", _pdf_bytes(""))
    assert att.text.strip() == ""
    assert "OCR" in att.note or "текстового слоя нет" in att.note


def test_cp1251_text_decoded(tmp_path):
    """Выгрузки из цеха часто в cp1251 — «кракозябр» быть не должно."""
    s = _session()
    att = af.attach_file(s, tmp_path, "note.txt",
                         "Мел, партия 12".encode("cp1251"))
    assert "Мел, партия 12" in att.text


# ----------------------------------------------------------------------
# 3. Отказы (A0.6 — с объяснением)
# ----------------------------------------------------------------------
def test_unsupported_format_rejected_with_list(tmp_path):
    s = _session()
    with pytest.raises(af.AttachmentError, match="не поддерживается"):
        af.attach_file(s, tmp_path, "photo.png", b"\x89PNG\r\n\x1a\n")


def test_empty_file_rejected(tmp_path):
    s = _session()
    with pytest.raises(af.AttachmentError, match="пуст"):
        af.attach_file(s, tmp_path, "empty.txt", b"")


def test_too_large_file_rejected(tmp_path):
    s = _session()
    with pytest.raises(af.AttachmentError, match="слишком большой"):
        af.attach_file(s, tmp_path, "big.txt", b"x" * 2048, max_bytes=1024)


def test_broken_xlsx_rejected_explicitly(tmp_path):
    s = _session()
    with pytest.raises(af.AttachmentError, match="Excel"):
        af.attach_file(s, tmp_path, "broken.xlsx", b"not really a workbook")


def test_failed_extraction_does_not_leave_file_or_record(tmp_path):
    """Отказ извлечения не оставляет мусора: ни записи, ни файла на диске."""
    s = _session()
    with pytest.raises(af.AttachmentError):
        af.attach_file(s, tmp_path, "broken.xlsx", b"garbage")
    assert s.attachments == []
    fdir = tmp_path / PROJECT / "assistant" / "files"
    assert not fdir.exists() or list(fdir.iterdir()) == []


# ----------------------------------------------------------------------
# 4. Усечение и дочитывание хвоста
# ----------------------------------------------------------------------
def test_long_text_truncated_in_session_but_full_on_disk(tmp_path):
    s = _session()
    payload = ("строка паспорта; " * 5000).encode("utf-8")
    att = af.attach_file(s, tmp_path, "long.txt", payload, max_text_chars=1000)

    assert att.truncated is True and len(att.text) == 1000
    assert att.n_chars > 1000 and "первые 1000" in att.note

    tail = af.attachment_text(s, tmp_path, "long.txt", start=1000, length=200)
    assert tail["total_chars"] == att.n_chars
    assert tail["has_more"] is True and len(tail["text"]) == 200


def test_attachment_text_reads_chunks_to_the_end(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "short.txt", b"abcdefghij")
    got = af.attachment_text(s, tmp_path, "short.txt", start=5, length=100)
    assert got["text"] == "fghij" and got["has_more"] is False


def test_attachment_text_unknown_name_lists_known(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "tds.txt", b"payload")
    with pytest.raises(af.AttachmentError, match="tds.txt"):
        af.attachment_text(s, tmp_path, "missing.txt")


def test_attachment_text_bad_range(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "tds.txt", b"payload")
    with pytest.raises(af.AttachmentError, match="length"):
        af.attachment_text(s, tmp_path, "tds.txt", length=0)


# ----------------------------------------------------------------------
# 5. Поиск, удаление, сироты
# ----------------------------------------------------------------------
def test_find_attachment_by_id_name_and_hash_prefix(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "TDS.pdf", _pdf_bytes("x"))
    assert af.find_attachment(s, att.id) is att
    assert af.find_attachment(s, "TDS.pdf") is att
    assert af.find_attachment(s, att.sha256[:10]) is att
    assert af.find_attachment(s, "нет такого") is None


def test_remove_attachment_deletes_file(tmp_path):
    s = _session()
    att = af.attach_file(s, tmp_path, "tds.txt", b"payload")
    path = af.attachment_path(tmp_path, PROJECT, att)
    assert af.remove_attachment(s, tmp_path, "tds.txt") is True
    assert s.attachments == [] and not path.exists()
    assert af.remove_attachment(s, tmp_path, "tds.txt") is False


def test_orphan_files_are_reported_not_deleted(tmp_path):
    """Файл без ссылки в сессии — ПОКАЗАТЬ, а не стереть молча (A0.6)."""
    s = _session()
    att = af.attach_file(s, tmp_path, "tds.txt", b"payload")
    s.remove_attachment(att.id)          # ссылку убрали, файл оставили
    orphans = af.orphan_files(s, tmp_path)
    assert len(orphans) == 1 and "tds" in orphans[0]


# ----------------------------------------------------------------------
# 6. Переживает сохранение/загрузку проекта и попадает в показ
# ----------------------------------------------------------------------
def test_attachments_survive_session_round_trip(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "TDS.pdf", _pdf_bytes("d50 = 2.5"))
    store.save_session(s, tmp_path)

    loaded = store.load_session(tmp_path, PROJECT)
    assert len(loaded.attachments) == 1
    got = af.attachment_text(loaded, tmp_path, "TDS.pdf")
    assert "d50" in got["text"]


def test_attachments_dataframe_and_digest(tmp_path):
    s = _session()
    af.attach_file(s, tmp_path, "lab.xlsx", _xlsx_bytes())
    df = views.attachments_dataframe(s)
    assert df.iloc[0]["файл"] == "lab.xlsx" and df.iloc[0]["символов"] > 0

    dig = views.attachment_digest(s, per_file_chars=50)
    assert dig[0]["name"] == "lab.xlsx" and len(dig[0]["text"]) <= 50
