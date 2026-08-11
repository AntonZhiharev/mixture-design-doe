"""assistant/files.py — вложения сессии: хранение, дедуп, извлечение текста (iter59).

Самый частый реальный вход технолога — документ: паспорт сырья (TDS, обычно
PDF), выгрузка лаборатории (xlsx/csv), протокол (docx), кусок спеки (json).
Ассистент обязан читать их, но при этом:

* **файл хранится в проекте** (``assistant/files/``) — разбор воспроизводим и
  переезжает вместе с кампанией;
* **дедуп по sha256** — один документ не занимает бюджет контекста дважды;
* **в контекст идёт ДАЙДЖЕСТ**, полный текст — только по явному запросу
  (:func:`attachment_text`), иначе один паспорт съест окно модели;
* **A0.6** — неподдерживаемый формат, битый файл, PDF-скан без текстового слоя
  дают ЯВНОЕ сообщение (``note``/исключение), а не тихо пустой текст. Пустое
  вложение выглядело бы как «в паспорте ничего нет» — худший вид молчания,
  потому что ассистенту запрещено выдумывать данные паспортов.

Зависимости уже в проекте: ``openpyxl`` (xlsx), ``python-docx`` (docx),
``pypdf`` (pdf) — добавлен в ``requirements.txt`` этим шагом.

iter68 — **изображения**. Скриншот экрана («вот такие границы вижу») и фото
паспорта — самый быстрый вход технолога, но текста в них нет: OCR мы не ставим,
читает их САМА модель (OpenRouter ``image_url``, см. :func:`data_url`).
Поэтому у такого вложения ``text`` пуст ОСОЗНАННО и это отмечено в ``note`` —
иначе пустой дайджест выглядел бы как «файл не прочитался» (A0.6).
"""
from __future__ import annotations

import base64
import hashlib
import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .session import AssistantSession, Attachment
from .store import files_dir

#: Максимальный размер вложения. Больше — явный отказ (документ такого размера
#: почти наверняка не паспорт, а выгрузка целиком; её место в базе, не в чате).
MAX_FILE_BYTES = 25 * 1024 * 1024

#: Сколько символов извлечённого текста хранить в сессии. Хвост остаётся на
#: диске и доступен через :func:`attachment_text` постранично.
MAX_TEXT_CHARS = 200_000

#: Ограничители табличных форматов — читаем «шапку и тело», а не гигабайты.
MAX_SHEET_ROWS = 500
MAX_SHEET_COLS = 60
MAX_PDF_PAGES = 200

_TEXT_EXT = {".txt", ".md", ".markdown", ".log", ".rst", ".ini", ".cfg",
             ".yaml", ".yml", ".csv", ".tsv", ".json", ".xml", ".html", ".py"}

_MIME_BY_EXT = {
    ".txt": "text/plain", ".md": "text/markdown", ".log": "text/plain",
    ".csv": "text/csv", ".tsv": "text/tab-separated-values",
    ".json": "application/json", ".xml": "application/xml",
    ".yaml": "application/yaml", ".yml": "application/yaml",
    ".html": "text/html", ".py": "text/x-python",
    ".xlsx": "application/vnd.openxmlformats-officedocument."
             "spreadsheetml.sheet",
    ".xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
    ".docx": "application/vnd.openxmlformats-officedocument."
             "wordprocessingml.document",
    ".pdf": "application/pdf",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".webp": "image/webp", ".gif": "image/gif",
}

#: Форматы изображений, которые принимает OpenRouter (`image_url`). Прочие
#: (bmp/tiff/heic) отклоняем ЯВНО: молча отправленный неподдерживаемый тип
#: вернулся бы невнятной ошибкой провайдера уже посреди хода.
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

#: Предел на КАРТИНКУ. Отдельный от :data:`MAX_FILE_BYTES`, потому что
#: изображение уходит в запрос модели целиком (base64 ≈ +33 %): 20-мегабайтный
#: скриншот — это отказ провайдера и сожжённый ход, а не «медленно».
MAX_IMAGE_BYTES = 8 * 1024 * 1024

#: Пометка вложения-изображения: у него нет текстового слоя ПО ЗАМЫСЛУ.
IMAGE_NOTE = ("изображение — текст не извлекается, картинку читает сама модель "
              "(нужна модель с поддержкой vision)")

#: Расширения, которые умеем читать. Всё прочее — явный отказ со списком.
SUPPORTED_EXT = sorted(set(_TEXT_EXT) | {".xlsx", ".xlsm", ".docx", ".pdf"}
                       | IMAGE_EXT)


class AttachmentError(ValueError):
    """Вложение принять нельзя — с объяснением причины (A0.6)."""


# ----------------------------------------------------------------------
# Вспомогательное
# ----------------------------------------------------------------------
def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def safe_filename(name: str) -> str:
    """Имя файла для диска: без путей и небезопасных символов.

    Исходное имя сохраняется в ``Attachment.name`` — пользователь видит своё,
    а на диске лежит предсказуемое.
    """
    base = os.path.basename(str(name or "").strip().replace("\\", "/"))
    base = re.sub(r"[^0-9A-Za-zА-Яа-яЁё._-]+", "_", base).strip("._") or "file"
    return base[:80]


def guess_mime(name: str) -> str:
    return _MIME_BY_EXT.get(Path(str(name)).suffix.lower(), "application/octet-stream")


def is_image_name(name: str) -> bool:
    """Похоже ли имя файла на изображение (по расширению)."""
    return Path(str(name or "")).suffix.lower() in IMAGE_EXT


def data_url(data: bytes, mime: str) -> str:
    """Байты картинки → ``data:<mime>;base64,…`` для OpenRouter ``image_url``.

    Провайдер принимает либо публичный URL, либо data-URL; у нас файл лежит
    локально в проекте, поэтому единственный рабочий путь — base64.
    """
    if not isinstance(data, (bytes, bytearray)) or not data:
        raise AttachmentError("Пустое содержимое изображения — нечего кодировать.")
    mime = str(mime or "").strip() or "image/png"
    return f"data:{mime};base64,{base64.b64encode(bytes(data)).decode('ascii')}"


def _decode(data: bytes) -> str:
    """Текст из байтов: utf-8 → cp1251 → utf-8 с заменой.

    Паспорта и выгрузки из цеха часто приходят в cp1251; молча получить
    «кракозябры» хуже, чем попробовать вторую кодировку.
    """
    for enc in ("utf-8", "cp1251"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# ----------------------------------------------------------------------
# Извлечение текста по форматам
# ----------------------------------------------------------------------
def _extract_json(data: bytes) -> Tuple[str, str]:
    raw = _decode(data)
    try:
        obj = json.loads(raw)
    except ValueError as exc:
        return raw, f"JSON не разобран ({exc}) — приложен как обычный текст."
    return json.dumps(obj, ensure_ascii=False, indent=2), ""


def _extract_xlsx(data: bytes) -> Tuple[str, str]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover — openpyxl в requirements
        raise AttachmentError("Для чтения .xlsx нужен пакет openpyxl.") from exc
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001 — формат книги бывает любым
        raise AttachmentError(f"Файл Excel не открылся: {exc}") from exc

    parts: List[str] = []
    notes: List[str] = []
    for ws in wb.worksheets:
        parts.append(f"### Лист: {ws.title}")
        n_rows = 0
        for row in ws.iter_rows(values_only=True):
            if n_rows >= MAX_SHEET_ROWS:
                notes.append(f"лист '{ws.title}': показаны первые "
                             f"{MAX_SHEET_ROWS} строк")
                break
            cells = ["" if v is None else str(v) for v in row[:MAX_SHEET_COLS]]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
            n_rows += 1
    wb.close()
    return "\n".join(parts), "; ".join(notes)


def _extract_docx(data: bytes) -> Tuple[str, str]:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover — python-docx в requirements
        raise AttachmentError("Для чтения .docx нужен пакет python-docx.") from exc
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise AttachmentError(f"Файл Word не открылся: {exc}") from exc

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for ti, table in enumerate(doc.tables, start=1):
        parts.append(f"### Таблица {ti}")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts), ""


def _extract_pdf(data: bytes) -> Tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover — pypdf в requirements
        raise AttachmentError(
            "Для чтения .pdf нужен пакет pypdf (pip install pypdf).") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise AttachmentError(f"PDF не открылся: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            raise AttachmentError(
                "PDF защищён паролем — снимите защиту и приложите снова.")

    parts: List[str] = []
    notes: List[str] = []
    pages = reader.pages
    if len(pages) > MAX_PDF_PAGES:
        notes.append(f"прочитаны первые {MAX_PDF_PAGES} из {len(pages)} страниц")
    for i, page in enumerate(pages[:MAX_PDF_PAGES], start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:  # noqa: BLE001 — битая страница не рушит документ
            txt = ""
            notes.append(f"страница {i} не прочиталась")
        if txt.strip():
            parts.append(f"— страница {i} —\n{txt.strip()}")

    text = "\n\n".join(parts)
    if not text.strip():
        # A0.6: пустой текст без объяснения = «в паспорте ничего нет»;
        # ассистенту запрещено выдумывать d50/плотности, поэтому причина
        # обязана быть видна человеку.
        notes.append("текстового слоя нет (похоже на скан) — нужен OCR или "
                     "текстовая версия паспорта")
    return text, "; ".join(notes)


def extract_text(name: str, data: bytes) -> Tuple[str, str]:
    """Извлечь текст из файла → ``(текст, примечание)``.

    Неподдерживаемое расширение — :class:`AttachmentError` со списком того,
    что принимаем (тихо приложить бинарник «как текст» = мусор в контексте).
    """
    ext = Path(str(name)).suffix.lower()
    if ext in IMAGE_EXT:
        # Текста нет ПО ЗАМЫСЛУ: картинку читает модель, а не мы. Пустая
        # строка без пояснения выглядела бы как «файл не прочитался».
        return "", IMAGE_NOTE
    if ext == ".json":
        return _extract_json(data)
    if ext in _TEXT_EXT:
        return _decode(data), ""
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(data)
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".pdf":
        return _extract_pdf(data)
    raise AttachmentError(
        f"Формат '{ext or '(без расширения)'}' не поддерживается. "
        f"Принимаются: {', '.join(SUPPORTED_EXT)}. Для .doc/.xls сохраните "
        f"в современном формате, для скана — приложите текстовую версию.")


# ----------------------------------------------------------------------
# Приложение файла к сессии
# ----------------------------------------------------------------------
def stored_name(sha256: str, name: str) -> str:
    """Имя файла на диске: ``<sha16>__<безопасное имя>`` (дедуп виден глазами)."""
    return f"{sha256[:16]}__{safe_filename(name)}"


def attachment_path(root: str | Path, project: str, att: Attachment) -> Path:
    return files_dir(root, project) / (att.stored_name
                                       or stored_name(att.sha256, att.name))


def attach_file(session: AssistantSession, root: str | Path, name: str,
                data: bytes, *, project: Optional[str] = None,
                max_bytes: int = MAX_FILE_BYTES,
                max_text_chars: int = MAX_TEXT_CHARS,
                note: str = "") -> Attachment:
    """Приложить файл к сессии: сохранить в проект + извлечь текст.

    Повторное приложение того же содержимого (совпал ``sha256``) НЕ создаёт
    вторую запись и не переписывает файл — возвращается существующее вложение
    (см. :meth:`AssistantSession.add_attachment`).
    """
    project = str(project or session.project or "").strip()
    if not project:
        raise AttachmentError(
            "Сессия не привязана к проекту — некуда сохранять файл. "
            "Сначала создайте/загрузите проект кампании.")
    if not isinstance(data, (bytes, bytearray)):
        raise AttachmentError("Ожидались байты содержимого файла.")
    data = bytes(data)
    if not data:
        raise AttachmentError(f"Файл '{name}' пуст (0 байт) — нечего читать.")
    if is_image_name(name) and len(data) > MAX_IMAGE_BYTES:
        # Отдельный предел: картинка уходит В ЗАПРОС целиком (base64 +33 %),
        # поэтому «слишком большая» здесь наступает раньше, чем для документа.
        raise AttachmentError(
            f"Изображение '{name}' слишком большое: {len(data) / 1048576:.1f} МБ "
            f"при лимите {MAX_IMAGE_BYTES / 1048576:.0f} МБ. Оно уходит в запрос "
            f"модели целиком — уменьшите масштаб или обрежьте до нужной части "
            f"экрана.")
    if len(data) > max_bytes:
        raise AttachmentError(
            f"Файл '{name}' слишком большой: {len(data) / 1048576:.1f} МБ при "
            f"лимите {max_bytes / 1048576:.0f} МБ. Приложите нужный фрагмент "
            f"или загрузите данные в базу опытов, а не в чат.")

    digest = sha256_bytes(data)
    existing = session.attachment_by_hash(digest)
    if existing is not None:
        return existing

    text, auto_note = extract_text(name, data)     # может бросить AttachmentError
    truncated = len(text) > max_text_chars
    n_chars = len(text)

    target_dir = files_dir(root, project)
    target_dir.mkdir(parents=True, exist_ok=True)
    fname = stored_name(digest, name)
    (target_dir / fname).write_bytes(data)

    notes = [n for n in (note, auto_note) if n]
    if truncated:
        notes.append(f"в контекст взяты первые {max_text_chars} символов из "
                     f"{n_chars} (полный текст — на диске)")

    att = Attachment(name=str(name), sha256=digest, size=len(data),
                     mime=guess_mime(name), stored_name=fname,
                     text=text[:max_text_chars], n_chars=n_chars,
                     truncated=truncated, note="; ".join(notes))
    return session.add_attachment(att)


def find_attachment(session: AssistantSession, ident: str
                    ) -> Optional[Attachment]:
    """Найти вложение по id, имени файла или префиксу sha256."""
    ident = str(ident or "").strip()
    if not ident:
        return None
    for a in session.attachments:
        if ident in (a.id, a.name, a.sha256):
            return a
    for a in session.attachments:
        if a.sha256.startswith(ident) or a.name.lower() == ident.lower():
            return a
    return None


def attachment_text(session: AssistantSession, root: str | Path, ident: str, *,
                    project: Optional[str] = None, start: int = 0,
                    length: int = 20_000) -> Dict[str, Any]:
    """Полный текст вложения ПОСТРАНИЧНО (инструмент чтения для ассистента).

    В контекст по умолчанию идёт дайджест; когда модели действительно нужен
    хвост паспорта, она запрашивает фрагмент отсюда. Возвращает словарь с
    текстом и позицией — по нему видно, есть ли продолжение.
    """
    project = str(project or session.project or "")
    att = find_attachment(session, ident)
    if att is None:
        known = ", ".join(a.name for a in session.attachments) or "(нет файлов)"
        raise AttachmentError(
            f"Вложение '{ident}' не найдено в сессии. Приложены: {known}.")
    if start < 0 or length <= 0:
        raise AttachmentError("start должен быть ≥ 0, length — > 0.")

    text = att.text
    if att.truncated:                    # хвост дочитываем с диска
        path = attachment_path(root, project, att)
        if path.exists():
            try:
                text, _ = extract_text(att.name, path.read_bytes())
            except AttachmentError:
                pass                     # остаёмся с усечённым текстом сессии
    chunk = text[start:start + length]
    return {"name": att.name, "sha256": att.sha256, "start": int(start),
            "length": len(chunk), "total_chars": len(text),
            "has_more": bool(start + len(chunk) < len(text)),
            "text": chunk, "note": att.note}


def attachment_data_url(session: AssistantSession, root: str | Path, ident: str,
                        *, project: Optional[str] = None) -> str:
    """Вложение-изображение → data-URL для запроса модели (iter68).

    Читаем с диска, а НЕ из сессии: base64 в ``session.json`` раздул бы файл
    переписки и бюджет контекста (оценка токенов считает символы), поэтому в
    сессии лежит только ссылка на файл, а картинка собирается на момент
    отправки.
    """
    project = str(project or session.project or "")
    att = find_attachment(session, ident)
    if att is None:
        known = ", ".join(a.name for a in session.attachments) or "(нет файлов)"
        raise AttachmentError(
            f"Вложение '{ident}' не найдено в сессии. Приложены: {known}.")
    if not is_image_name(att.name):
        raise AttachmentError(
            f"Вложение '{att.name}' не изображение ({att.mime or 'тип неизвестен'}): "
            f"как картинку его отправить нельзя. Текстовые документы модель "
            f"читает инструментом read_attachment.")
    path = attachment_path(root, project, att)
    if not path.exists():
        raise AttachmentError(
            f"Файл изображения '{att.name}' не найден на диске ({path}). "
            f"Возможно, проект переносили без каталога assistant/files.")
    return data_url(path.read_bytes(), att.mime or guess_mime(att.name))


def remove_attachment(session: AssistantSession, root: str | Path, ident: str,
                      *, project: Optional[str] = None,
                      delete_file: bool = True) -> bool:
    """Убрать вложение из сессии (по умолчанию — и файл с диска)."""
    project = str(project or session.project or "")
    att = find_attachment(session, ident)
    if att is None:
        return False
    ok = session.remove_attachment(att.id)
    if ok and delete_file:
        path = attachment_path(root, project, att)
        try:
            if path.exists():
                path.unlink()
        except OSError:
            pass                          # файл держит другой процесс — не беда
    return ok


def orphan_files(session: AssistantSession, root: str | Path, *,
                 project: Optional[str] = None) -> List[str]:
    """Файлы в ``assistant/files``, на которые сессия уже не ссылается.

    Диагностика (A0.6): показать, а не удалять молча — файл мог остаться от
    сессии, очищенной пользователем, и всё ещё быть нужен как документ.
    """
    project = str(project or session.project or "")
    d = files_dir(root, project)
    if not d.exists():
        return []
    known = {a.stored_name or stored_name(a.sha256, a.name)
             for a in session.attachments}
    return sorted(p.name for p in d.iterdir()
                  if p.is_file() and p.name not in known)
